'''
	Python script to count words in a Microsoft Word document.
	Will accept command line an input file and optionally an output file.
'''

import sys, re, os, argparse
from docx import Document
import subprocess as sub

def load_document(filename):
	document = Document(filename)
	return document

def extract_words(doc):
	para_set = set()
	for i in range(len(doc.paragraphs)):
		para_set.add(doc.paragraphs[i].text)
	word_list = []
	for i in para_set:
		words = re.split('\s+', i)
		for p, word in enumerate(words):
			word = re.sub(u'\u201c','"', word)
			word = re.sub(u'\u201d', '"', word)
			word = word.lower().strip(",.\?\"\!")
			word_list.append(word)
	word_set = set(word_list)
	return word_list, word_set

def get_word_count(w_list, w_set):
	word_count = {}
	for word in w_set:
		counter = 0
		for i in range(len(w_list)):
			if word == w_list[i]:
				counter += 1
		word_count[word] = counter
	return sorted(word_count.items(), key=lambda kv: kv[1], reverse=True)

if __name__ == '__main__':
	parser = argparse.ArgumentParser(description='Counts the words in a Microsoft Word document.')
	parser.add_argument('input_file', action='store', type=str, help='Enter the name of the file you want to perform the word count on.')
	group = parser.add_mutually_exclusive_group(required=False)
	group.add_argument('-o', '--output', action='store', default=False, help='Enter filename for output')
	group.add_argument('-v', '--verbose', action='store_true', default=True, help='To print results to screen')
	args = parser.parse_args()
		
	file_name = args.input_file
	in_doc = load_document(file_name)
	list_doc, set_doc = extract_words(in_doc)
	count_doc = get_word_count(list_doc, set_doc)
	if args.output:
		with open(args.output, 'w') as f:
			for i, line in enumerate(count_doc):
				f.write('%d %s : %d\n' % (i+1, line[0], line[1]))
	if args.verbose:
		for i, line in enumerate(count_doc):
			if line[1] >= 3:
				print('%d %s : %d' % (i+1, line[0], line[1]))
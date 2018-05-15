'''
	Python script to count words in a Microsoft Word document.
	Will accept command line an input file and optionally an output file.
'''

import sys, re, os, argparse
from docx import Document
import subprocess as sub

def load_document(filename):
	document = Document(filename)
	return document

def extract_sentences(doc):
	para_set = set()
	for i in range(len(doc.paragraphs)):
		para_set.add(doc.paragraphs[i].text)
	sentence_list = []
	for i in para_set:
		sentences = re.split('\b', i)
		for p, sentence in enumerate(sentences):
			sentence = re.sub(u'\u201c','"', sentence)
			sentence = re.sub(u'\u201d', '"', sentence)
			sentence = sentence.lower().strip(",.\?\"\!")
			sentence_list.append(sentence)
	sentence_set = set(sentence_list)
	return sentence_list, sentence_set

def get_sentence_count(w_list, w_set):
	sentence_count = {}
	for sentence in w_set:
		counter = 0
		for i in range(len(w_list)):
			if sentence == w_list[i]:
				counter += 1
		sentence_count[sentence] = counter
	return sorted(sentence_count.items(), key=lambda kv: kv[1], reverse=True)

if __name__ == '__main__':
	parser = argparse.ArgumentParser(description='Counts the words in a Microsoft Word document.')
	parser.add_argument('input_file', action='store', type=str, help='Enter the name of the file you want to perform the word count on.')
	group = parser.add_mutually_exclusive_group(required=False)
	group.add_argument('-o', '--output', action='store', default=False, help='Enter filename for output')
	group.add_argument('-v', '--verbose', action='store_true', default=True, help='To print results to screen')
	args = parser.parse_args()
		
	file_name = args.input_file
	in_doc = load_document(file_name)
	list_doc, set_doc = extract_sentences(in_doc)
	count_doc = get_sentence_count(list_doc, set_doc)
	if args.output:
		with open(args.output, 'w') as f:
			for i, line in enumerate(count_doc):
				f.write('%d %s : %d\n' % (i+1, line[0], line[1]))
	if args.verbose:
		for i, line in enumerate(count_doc):
			if line[1] >= 3:
				print('%d %s : %d' % (i+1, line[0], line[1]))